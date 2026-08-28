import tkinter as tk
from tkinter import filedialog, messagebox
import os
import copy
from docx import Document

def clean_and_sort_word_doc():
    # 1. Hide the main blank tkinter window
    root = tk.Tk()
    root.withdraw()

    print("Please select your Word (.docx) file from the popup window...")

    # 2. Open file selection dialog
    input_path = filedialog.askopenfilename(
        title="Select your References Word File",
        filetypes=[("Word Documents", "*.docx")]
    )

    if not input_path:
        print("No file selected. Exiting.")
        return

    try:
        doc = Document(input_path)
    except Exception as e:
        messagebox.showerror("Error", f"Could not open the file.\nError: {e}")
        return

    # 3. Read paragraphs and preserve formatting (XML)
    #    We store them as a list of tuples: (Plain_Text, Full_XML_Element)
    content_list = []
    
    for paragraph in doc.paragraphs:
        text_content = paragraph.text.strip()
        
        # Only keep lines that have text (ignore empty blank lines for now)
        if text_content:
            content_list.append((text_content, paragraph._element))

    total_original = len(content_list)

    # 4. Remove Duplicates
    #    We check if the 'text' is already seen.
    seen_text = set()
    unique_items = []
    
    for text, element in content_list:
        if text not in seen_text:
            # Create a deep copy of the element so we don't mess up the original references
            unique_items.append((text, copy.deepcopy(element)))
            seen_text.add(text)
    
    # 5. Sort Alphabetically
    #    Sort based on the plain text (Case insensitive)
    unique_items.sort(key=lambda x: x[0].lower())

    total_final = len(unique_items)
    duplicates_removed = total_original - total_final

    # 6. Rebuild the Document
    #    We clear the document body and add the sorted elements back in.
    body = doc._element.body
    body.clear_content()

    for text, element in unique_items:
        body.append(element)
        
        # OPTIONAL: If you want a specific manual blank line between every reference,
        # uncomment the two lines below. (Usually Word handles this with 'Space After' paragraph style)
        # p = doc.add_paragraph()
        # p.paragraph_format.space_after = 0

    # 7. Generate Output Filename
    folder = os.path.dirname(input_path)
    filename = os.path.basename(input_path)
    name_part, ext_part = os.path.splitext(filename)
    
    output_filename = f"{name_part}_organized{ext_part}"
    output_path = os.path.join(folder, output_filename)

    # 8. Save
    try:
        doc.save(output_path)
        
        success_msg = (
            f"Success!\n\n"
            f"Original paragraphs: {total_original}\n"
            f"Duplicates removed: {duplicates_removed}\n"
            f"Final sorted count: {total_final}\n\n"
            f"Saved as:\n{output_filename}"
        )
        print(success_msg)
        messagebox.showinfo("Done", success_msg)
        
    except Exception as e:
        messagebox.showerror("Error", f"Could not save the file.\nMake sure the file isn't open in Word!\n\nError: {e}")

if __name__ == "__main__":
    clean_and_sort_word_doc()
